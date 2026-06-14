"""Generate the binary test fixtures (sample.pdf, sample.docx).

These files are committed to the repo so tests don't depend on this script at
runtime. Regenerate them with:

    python tests/fixtures/generate_fixtures.py

Requires the dev dependencies (``fpdf2`` and ``python-docx``).
"""

from pathlib import Path

from docx import Document as DocxDocument
from fpdf import FPDF

FIXTURES_DIR = Path(__file__).parent


def generate_pdf() -> None:
    """A two-page PDF, with a distinct marker on each page."""
    pdf = FPDF()
    pdf.set_font("Helvetica", size=12)

    pdf.add_page()
    pdf.multi_cell(
        0,
        10,
        "Chapter One\n\n"
        "This is the first page of the sample PDF. It contains the marker "
        "PdfPageOne so the parser test can verify page one is extracted with "
        "page_number == 1.",
    )

    pdf.add_page()
    pdf.multi_cell(
        0,
        10,
        "Chapter Two\n\n"
        "This is the second page of the sample PDF. It contains the marker "
        "PdfPageTwo so the parser test can verify page two is extracted with "
        "page_number == 2.",
    )

    pdf.output(str(FIXTURES_DIR / "sample.pdf"))


def generate_docx() -> None:
    """A DOCX with a Heading 1 / Heading 2 outline and body paragraphs."""
    document = DocxDocument()

    document.add_heading("Chapter 1", level=1)
    document.add_paragraph("Introduction paragraph for the first chapter (DocxIntro).")

    document.add_heading("Section 1.1", level=2)
    document.add_paragraph("Content of section one point one containing DocxSection11.")

    document.add_heading("Chapter 2", level=1)
    document.add_heading("Section 2.1", level=2)
    document.add_paragraph(
        "Content of section two point one containing the unique marker DocxSection21, "
        "with enough text to represent a realistic paragraph in the document outline."
    )

    document.save(str(FIXTURES_DIR / "sample.docx"))


if __name__ == "__main__":
    generate_pdf()
    generate_docx()
    print(f"Wrote fixtures to {FIXTURES_DIR}")
